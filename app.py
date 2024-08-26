import streamlit as st
import requests
import json
from docx import Document
from io import BytesIO

# Set page configuration
st.set_page_config(page_title="Diccionario de Problemas Económicos", page_icon="📚", layout="wide")

# Function to set custom CSS for a more elegant design
def set_custom_css():
    st.markdown(
        """
        <style>
        .stApp {
            background-color: #FFF9C4; /* Light yellow background color */
            font-family: 'Roboto', sans-serif;
        }
        .title {
            font-size: 2.5em;
            color: #333;
            text-align: center;
            margin-bottom: 1rem;
        }
        .content-box {
            background: #ffffffcc;
            padding: 2rem;
            border-radius: 10px;
            box-shadow: 0 4px 6px rgba(0, 0, 0, 0.1);
        }
        .info-box {
            background: #4CAF50;
            color: white;
            padding: 1rem 1.5rem;
            border-radius: 8px;
            margin-bottom: 1rem;
        }
        .stButton > button {
            background: #4CAF50;
            color: white;
            border-radius: 8px;
            padding: 0.75rem 1.5rem;
            border: none;
            cursor: pointer;
            transition: background-color 0.3s ease;
        }
        .stButton > button:hover {
            background-color: #45a049;
        }
        </style>
        """,
        unsafe_allow_html=True
    )

# Function to create the information column
def crear_columna_info():
    st.markdown("""
    <div class="info-box">
    ## Sobre esta aplicación

    Esta aplicación es un Diccionario de Problemas Económicos. Permite a los usuarios obtener respuestas a problemas económicos según la interpretación de diversas corrientes económicas.

    ### Cómo usar la aplicación:

    1. Elija un problema económico de la lista predefinida o proponga su propio problema.
    2. Seleccione una o más corrientes económicas.
    3. Haga clic en "Obtener respuesta" para generar las respuestas.
    4. Lea las respuestas y fuentes proporcionadas.
    5. Si lo desea, descargue un documento DOCX con toda la información.

    ### Autor y actualización:
    **Moris Polanco**, 26 ag 2024

    ### Cómo citar esta aplicación (formato APA):
    Polanco, M. (2024). *Diccionario de Problemas Económicos* [Aplicación web]. https://dicproblemas.streamlit.app

    ---
    **Nota:** Esta aplicación utiliza inteligencia artificial para generar respuestas basadas en información disponible en línea. Siempre verifique la información con fuentes académicas para un análisis más profundo.
    </div>
    """, unsafe_allow_html=True)

# Titles and Main Column
st.markdown("<div class='title'>Diccionario de Problemas Económicos</div>", unsafe_allow_html=True)

# Apply custom CSS
set_custom_css()

col1, col2 = st.columns([1, 2])

with col1:
    crear_columna_info()

with col2:
    TOGETHER_API_KEY = st.secrets["TOGETHER_API_KEY"]
    SERPER_API_KEY = st.secrets["SERPER_API_KEY"]

    # 101 economic problems in question form
    problemas_economicos = sorted([
        "¿Cómo se determinan los precios en un mercado libre?", "¿Cuál es el impacto del salario mínimo en el empleo?", 
        "¿Qué efectos tiene la inflación en la economía?", "¿Cómo influyen los tipos de interés en la inversión?", 
        "¿Es el déficit fiscal perjudicial para la economía?", "¿Cómo afectan los aranceles al comercio internacional?", 
        "¿Qué papel juegan los bancos centrales en la estabilidad económica?", 
        "¿La política fiscal expansionista es efectiva para combatir la recesión?", 
        "¿Cuál es la relación entre la oferta monetaria y la inflación?", 
        "¿Qué diferencias existen entre el capitalismo y el socialismo?", "¿Cómo funciona el concepto de la mano invisible de Adam Smith?", 
        "¿Cuáles son las causas del desempleo estructural?", "¿Qué ventaja tiene el comercio internacional?", 
        "¿Cómo afecta la globalización a las economías locales?", "¿Qué es la teoría del valor-trabajo?", 
        "¿Cuáles son los problemas del mercado laboral contemporáneo?", "¿De qué manera impacta la educación en el crecimiento económico?", 
        "¿Qué es la trampa de la liquidez en la teoría keynesiana?", "¿Puede el proteccionismo ser beneficioso?", 
        "¿Qué son los ciclos económicos?", "¿Cómo se mide el producto interno bruto (PIB)?", 
        "¿Qué es la estanflación y cómo se combate?", "¿Cómo se relacionan la oferta y demanda de dinero con la inflación?", 
        "¿Cuál es el efecto de la deuda pública sobre la economía?", "¿Qué es la teoría de los juegos y cómo se aplica en la economía?", 
        "¿Cómo afectan los impuestos al comportamiento de los consumidores?", "¿Qué es la elasticidad precio de la demanda?", 
        "¿Qué implica la existencia de externalidades en el mercado?", "¿Cómo se determina el tipo de cambio en el mercado de divisas?", 
        "¿Qué es la curva de Laffer?", "¿Qué es el análisis coste-beneficio?", "¿Qué papel juegan las instituciones en el desarrollo económico?", 
        "¿Cómo se regula la competencia en los mercados?", "¿Qué es la economía del bienestar?", 
        "¿Cuáles son las consecuencias de la intervención gubernamental en la economía?", "¿Qué es la eficiencia de Pareto?", 
        "¿Qué es la teoría del capital humano?", "¿Qué es el coeficiente de Gini y qué mide?", 
        "¿Cómo funciona el modelo de oferta y demanda agregada?", "¿Qué es la teoría económica neoclásica?", 
        "¿Qué es el equilibrio general en la economía?", "¿Cómo afectan las políticas monetarias a la economía?", 
        "¿Qué es la matriz insumo-producto?", "¿Qué es la paradoja del ahorro?", "¿Qué es la teoría de las expectativas racionales?", 
        "¿Qué es el multiplicador keynesiano?", "¿Qué es un monopolio y cómo se regula?", "¿Qué es la teoría de la utilidad marginal?", 
        "¿Qué es la teoría de la renta ricardiana?", "¿Qué es la teoría de la elección del consumidor?", 
        "¿Qué es el equilibrio de Nash?", "¿Qué es la teoría del comercio internacional?", 
        "¿Qué es la economía de la oferta?", "¿Qué son las economías de escala?", 
        "¿Qué es la teoría de la agencia?", "¿Qué es la teoría del desarrollo endógeno?", 
        "¿Qué es el superávit y el déficit comercial?", "¿Qué son los bienes públicos y cómo se financian?", 
        "¿Qué es la economía conductual?", "¿Qué es la teoría del crecimiento económico?", 
        "¿Qué es la tasa natural de desempleo?", "¿Qué son las políticas de estabilización?", 
        "¿Qué es la teoría de la preferencia intertemporal?", "¿Qué es la regla de oro de la acumulación?", 
        "¿Para qué sirve el modelo de Heckscher-Ohlin?", "¿Qué es la política monetaria no convencional?", 
        "¿Qué es la economía de la información?", "¿Cómo afectan los subsidios a la economía?", 
        "¿Qué es la teoría del salario de eficiencia?", "¿Qué es el efecto multiplicador del gasto público?", 
        "¿Qué es la política de ingresos?", "¿Qué es el mercado laboral dual?", 
        "¿Qué es el modelo de crecimiento de Solow?", "¿Qué es la teoría de los mercados eficientes?", 
        "¿Qué es el crowding out?", "¿Qué es la economía del comportamiento?", "¿Qué es la economía informal?", 
        "¿Qué es el problema de agencia?", "¿Qué es la teoría del estado estacionario?", 
        "¿Qué es la economía institucional?", "¿Cómo afectan las expectativas en los mercados financieros?", 
        "¿Qué es el ciclo económico real?", "¿Qué es el mercado de capitales?", 
        "¿Qué es la teoría cuantitativa del dinero?", "¿Qué es la inflación por demanda?", 
        "¿Qué es la inflación por costos?", "¿Qué es la política arancelaria?", 
        "¿Qué es la paridad del poder adquisitivo?", "¿Qué es la solvencia fiscal?", 
        "¿Qué son los bienes de lujo?", "¿Qué son los bienes inferiores?", "¿Qué son los bienes sustitutos?", 
        "¿Qué es la curva de demanda?", "¿Qué es la oferta agregada?", "¿Qué es el producto potencial?", 
        "¿Qué es el modelo IS-LM?", "¿Qué es la política de oferta?", "¿Qué es el balance presupuestario?", 
        "¿Qué son los bienes transables?", "¿Qué es la economía política?", "¿Qué es la economía experimental?", 
        "¿Qué es la teoría de los contratos?"
    ])

    # Economic schools of thought
    escuelas_economicas = [
        "Escuela Austríaca de Economía", "Socialismo", "Keynesianismo", "Monetarismo", 
        "Mercantilismo", "Marxismo", "Neoclasicismo", "Fisiocracia", "Economía del Desarrollo", 
        "Economía Conductual", "Escuela de Chicago", "Institucionalismo", "Económica Feminista", 
        "Poskeynesianismo", "Escuela de Fráncfort"
    ]

    def buscar_informacion(query, escuela):
        url = "https://google.serper.dev/search"
        payload = json.dumps({
            "q": f"{query} {escuela} economía"
        })
        headers = {
            'X-API-KEY': SERPER_API_KEY,
            'Content-Type': 'application/json'
        }
        response = requests.post(url, headers=headers, data=payload)
        return response.json()

    def generar_respuesta(problema, escuela, contexto):
        url = "https://api.together.xyz/inference"
        payload = json.dumps({
            "model": "mistralai/Mixtral-8x7B-Instruct-v0.1",
            "prompt": f"Contexto: {contexto}\n\nProblema: {problema}\nEscuela: {escuela}\n\nProporciona una respuesta al problema económico '{problema}' según la interpretación del {escuela}. La respuesta debe ser concisa pero informativa, similar a una entrada de diccionario. Si es posible, incluye una referencia a una obra o figura específica de {escuela} que trate este concepto.\n\nRespuesta:",
            "max_tokens": 2048,
            "temperature": 0,
            "top_p": 0.7,
            "top_k": 50,
            "repetition_penalty": 0,
            "stop": ["Problema:"]
        })
        headers = {
            'Authorization': f'Bearer {TOGETHER_API_KEY}',
            'Content-Type': 'application/json'
        }
        response = requests.post(url, headers=headers, data=payload)
        return response.json()['output']['choices'][0]['text'].strip()

    def create_docx(problema, respuestas, fuentes):
        doc = Document()
        doc.add_heading('Diccionario de Problemas Económicos', 0)

        doc.add_heading('Problema', level=1)
        doc.add_paragraph(problema)

        for escuela, respuesta in respuestas.items():
            doc.add_heading(f'Respuesta según la corriente {escuela}', level=2)
            doc.add_paragraph(respuesta)

        doc.add_heading('Fuentes', level=1)
        for fuente in fuentes:
            doc.add_paragraph(fuente, style='List Bullet')

        doc.add_paragraph('\nNota: Este documento fue generado por un asistente de IA. Verifica la información con fuentes académicas para un análisis más profundo.')

        return doc

    with st.container():
        st.markdown("<div class='content-box'>", unsafe_allow_html=True)
        st.write("Elige un problema económico de la lista o propón tu propio problema:")

        opcion = st.radio("", ["Elegir de la lista", "Proponer mi propio problema"])

        if opcion == "Elegir de la lista":
            problema = st.selectbox("Selecciona un problema:", problemas_economicos)
        else:
            problema = st.text_input("Ingresa tu propio problema económico:")

        st.write("Selecciona una o más corrientes económicas (máximo 5):")
        escuelas_seleccionadas = st.multiselect("Corrientes Económicas", escuelas_economicas)

        if len(escuelas_seleccionadas) > 5:
            st.warning("Has seleccionado más de 5 corrientes. Por favor, selecciona un máximo de 5.")
        else:
            if st.button("Obtener respuesta"):
                if problema and escuelas_seleccionadas:
                    with st.spinner("Buscando información y generando respuestas..."):
                        respuestas, todas_fuentes = {}, []

                        for escuela in escuelas_seleccionadas:
                            # Buscar información relevante
                            resultados_busqueda = buscar_informacion(problema, escuela)
                            contexto = "\n".join([item["snippet"] for item in resultados_busqueda.get("organic", [])])
                            fuentes = [item["link"] for item in resultados_busqueda.get("organic", [])]

                            # Generar respuesta
                            respuesta = generar_respuesta(problema, escuela, contexto)

                            respuestas[escuela] = respuesta
                            todas_fuentes.extend(fuentes)

                        # Mostrar las respuestas
                        st.subheader(f"Respuestas para el problema: {problema}")
                        for escuela, respuesta in respuestas.items():
                            st.markdown(f"**{escuela}:** {respuesta}")

                        # Botón para descargar el documento
                        doc = create_docx(problema, respuestas, todas_fuentes)
                        buffer = BytesIO()
                        doc.save(buffer)
                        buffer.seek(0)
                        st.download_button(
                            label="Descargar respuesta en DOCX",
                            data=buffer,
                            file_name=f"Respuesta_{problema.replace(' ', '_')}.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                        )
                else:
                    st.warning("Por favor, selecciona un problema y al menos una corriente.")
        st.markdown("</div>", unsafe_allow_html=True)
